"""Generate the thesis plan as a Word .docx file."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    pass


def add_heading_styled(doc, text, level=1, color=None, size=None):
    h = doc.add_heading(text, level=level)
    if color or size:
        for run in h.runs:
            if color:
                run.font.color.rgb = color
            if size:
                run.font.size = size
    return h


def p(doc, text, bold=False, italic=False, size=None, align=None, indent=None):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    if indent is not None:
        para.paragraph_format.left_indent = Cm(indent)
    run = para.add_run(text)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = size
    return para


def p_mixed(doc, parts, indent=None, align=None):
    """parts = list of tuples (text, bold, italic)."""
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    if indent is not None:
        para.paragraph_format.left_indent = Cm(indent)
    for text, bold, italic in parts:
        run = para.add_run(text)
        run.font.name = "Calibri"
        run.bold = bold
        run.italic = italic
    return para


def bullet(doc, text, level=0, italic=False):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Cm(0.75 + level * 0.75)
    run = para.runs[0] if para.runs else para.add_run("")
    if not para.runs:
        run = para.add_run(text)
    else:
        run.text = text
    run.font.name = "Calibri"
    run.italic = italic
    return para


def ref(doc, text, indent=1.5):
    """Add an APA-style reference in italics with arrow prefix."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(indent)
    arrow = para.add_run("→ ")
    arrow.font.name = "Calibri"
    r = para.add_run(text)
    r.font.name = "Calibri"
    r.italic = True
    return para


def separator(doc, char="═", n=60):
    p(doc, char * n, align=WD_ALIGN_PARAGRAPH.CENTER)


def main():
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============================================================
    # COVER / TITLE
    # ============================================================
    p(doc, "PROJET DE FIN D'ÉTUDES — MASTER FINANCE ET DATA SCIENCE",
      bold=True, size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "FSJES Mohammedia — Université Hassan II de Casablanca",
      italic=True, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "En collaboration avec PEAQOCK", italic=True,
      size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "")
    p(doc, "PLAN DÉTAILLÉ DU MÉMOIRE", bold=True, size=Pt(16),
      align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "")

    separator(doc)
    p(doc, "TITRE DU MÉMOIRE", bold=True, size=Pt(14),
      align=WD_ALIGN_PARAGRAPH.CENTER)
    separator(doc)

    p_mixed(doc, [
        ("« Conception et implémentation d'un agent multi-signaux d'aide "
         "à la décision d'investissement pour les marchés émergents : "
         "application à l'action Attijariwafa Bank à la Bourse de "
         "Casablanca »", True, False)
    ], align=WD_ALIGN_PARAGRAPH.CENTER)

    p_mixed(doc, [
        ("Sous-titre académique : ", False, True),
        ("Une approche hybride combinant analyse fondamentale, analyse "
         "technique, indicateurs macroéconomiques et traitement automatique "
         "du langage naturel des actualités financières.", False, False)
    ], align=WD_ALIGN_PARAGRAPH.CENTER)

    p(doc, "")
    p(doc, "Étudiant : Omar ARHOUNE", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "Encadrante académique : Mme Amal EL MZABI",
      align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "Entreprise d'accueil : Peaqock",
      align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ============================================================
    # PROBLÉMATIQUE
    # ============================================================
    separator(doc)
    p(doc, "PROBLÉMATIQUE CENTRALE", bold=True, size=Pt(14),
      align=WD_ALIGN_PARAGRAPH.CENTER)
    separator(doc)

    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Cm(1)
    quote.paragraph_format.right_indent = Cm(1)
    qrun = quote.add_run(
        "« Dans quelle mesure un agent algorithmique multi-signaux à base "
        "de règles pondérées, intégrant simultanément des dimensions "
        "fondamentale, technique, macroéconomique et médiatique, peut-il "
        "améliorer la qualité de la décision d'investissement sur un "
        "marché émergent peu liquide tel que la Bourse de Casablanca, "
        "en particulier sur le titre Attijariwafa Bank ? »")
    qrun.bold = True
    qrun.italic = True

    p(doc, "")
    p(doc, "Sous-questions de recherche :", bold=True)

    sub_questions = [
        "Quels sont les apports et les limites des théories financières "
        "classiques (efficience des marchés, finance comportementale, "
        "valorisation fondamentale, analyse technique) lorsqu'elles sont "
        "confrontées aux spécificités structurelles d'un marché émergent à "
        "faible liquidité comme la Bourse de Casablanca ?",
        "Quelle architecture technique et quelle logique de pondération "
        "permettent d'agréger de manière transparente, auditable et "
        "reproductible des signaux hétérogènes (quantitatifs, qualitatifs, "
        "médiatiques) en une recommandation d'investissement actionnable ?",
        "L'intégration d'une couche d'analyse de sentiment fondée sur des "
        "sources d'information marocaines apporte-t-elle un gain "
        "informationnel mesurable par rapport à un système purement "
        "quantitatif ?",
        "Au regard d'un backtesting rigoureux (rendement, ratio de Sharpe, "
        "drawdown maximal) sur un historique pertinent, l'agent "
        "multi-signaux surperforme-t-il une stratégie passive de type "
        "« buy-and-hold » sur l'indice MASI ou sur le titre "
        "Attijariwafa Bank ?",
    ]
    for i, q in enumerate(sub_questions, 1):
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(q)
        run.font.name = "Calibri"

    doc.add_page_break()

    # ============================================================
    # PAGES PRÉLIMINAIRES
    # ============================================================
    separator(doc)
    p(doc, "PLAN DÉTAILLÉ DU MÉMOIRE (80–120 pages)",
      bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)
    separator(doc)

    add_heading_styled(doc, "PAGES PRÉLIMINAIRES", level=1)
    prelim = [
        "Page de garde",
        "Résumé en français (≈ 250 mots)",
        "Abstract en anglais",
        "ملخص بالعربية",
        "Dédicace",
        "Remerciements (Mme Amal EL MZABI, Peaqock, FSJES Mohammedia – "
        "Université Hassan II de Casablanca, famille)",
        "Sommaire",
        "Liste des abréviations, sigles et acronymes (ATW, MASI, AMMC, "
        "BVC, DCF, DDM, EMH, NLP, RSI, MACD, VaR, CAPM, MEDAF, GBM, etc.)",
        "Liste des tableaux",
        "Liste des figures",
    ]
    for item in prelim:
        doc.add_paragraph(item, style="List Bullet")

    # ============================================================
    # INTRODUCTION GÉNÉRALE
    # ============================================================
    add_heading_styled(doc, "INTRODUCTION GÉNÉRALE (pages 1–8)", level=1)

    add_heading_styled(doc, "1. Contexte et mise en situation", level=2)
    p(doc, "Trois paragraphes introductifs ancrés dans la littérature :",
      italic=True)

    p_mixed(doc, [
        ("Paragraphe 1 — La révolution algorithmique de la finance. ",
         True, True),
        ("Mobiliser : Aldridge (2013) sur le trading haute fréquence, "
         "Lopez de Prado (2018) Advances in Financial Machine Learning, "
         "Hendershott, Jones & Menkveld (2011) sur l'impact de "
         "l'algorithmique sur la liquidité.", False, False)
    ])

    p_mixed(doc, [
        ("Paragraphe 2 — Les marchés émergents et leurs spécificités. ",
         True, True),
        ("Mobiliser : Bekaert & Harvey (1997, 2002), Lesmond (2005) sur "
         "la liquidité, Harvey (1995) sur la prévisibilité des rendements "
         "émergents.", False, False)
    ])

    p_mixed(doc, [
        ("Paragraphe 3 — La Bourse de Casablanca et l'enjeu national. ",
         True, True),
        ("Citer rapports AMMC, BVC, Bank Al-Maghrib, et la position "
         "systémique d'Attijariwafa Bank dans l'indice MASI.", False, False)
    ])

    add_heading_styled(doc, "2. Objectif du travail", level=2)
    p(doc, "Concevoir, implémenter et évaluer un agent algorithmique "
      "d'aide à la décision combinant quatre familles de signaux pour "
      "produire un verdict (ACHAT / CONSERVER / VENDRE) explicable et "
      "auditable sur le titre ATW.")

    add_heading_styled(doc, "3. Problématique et sous-questions", level=2)
    p(doc, "(Rappel — voir ci-dessus.)", italic=True)

    add_heading_styled(doc, "4. Méthodologie", level=2)
    method_items = [
        ("Méthode de recherche : ", "approche mixte (revue systématique + "
         "démarche design science research)."),
        ("Sources documentaires : ", "journaux académiques classés (FT50, "
         "ABS), littérature grise (rapports AMMC, BVC, BAM, FMI, Banque "
         "Mondiale)."),
        ("Données primaires : ", "scraping de 6 sources de presse "
         "marocaines, données BVC, Medias24, MarketScreener."),
        ("Outils : ", "Python 3.10+, PostgreSQL, Docker, Selenium, "
         "scikit-learn, NLTK/spaCy, Groq LLM."),
        ("Démarche : ", "(a) revue théorique → (b) modélisation "
         "conceptuelle → (c) implémentation → (d) backtesting → "
         "(e) interprétation."),
    ]
    for label, txt in method_items:
        para = doc.add_paragraph(style="List Bullet")
        run1 = para.add_run(label)
        run1.bold = True
        run1.font.name = "Calibri"
        run2 = para.add_run(txt)
        run2.font.name = "Calibri"

    add_heading_styled(doc, "5. Structure du mémoire", level=2)
    p(doc, "Annonce des trois chapitres : arsenal théorique, contexte "
      "d'application, conception et implémentation pratique.")

    doc.add_page_break()

    # ============================================================
    # CHAPITRE 1
    # ============================================================
    add_heading_styled(
        doc,
        "CHAPITRE 1 — ARSENAL CONCEPTUEL ET THÉORIQUE : DES THÉORIES "
        "DE MARCHÉ AUX SYSTÈMES ALGORITHMIQUES MULTI-SIGNAUX (pages 9–45)",
        level=1)

    p_mixed(doc, [("Introduction du Chapitre 1", True, True)])

    # Section 1
    add_heading_styled(doc,
                       "Section 1 — Fondements théoriques de la décision "
                       "d'investissement en bourse", level=2)

    add_heading_styled(doc, "I. La théorie de l'efficience des marchés et "
                       "ses contestations", level=3)

    p_mixed(doc, [("A. L'hypothèse d'efficience informationnelle (EMH)",
                   True, False)])
    p(doc, "1. Les trois formes d'efficience : faible, semi-forte, forte",
      indent=1)
    ref(doc, "Fama, E. F. (1970). Efficient capital markets: A review of "
        "theory and empirical work. The Journal of Finance, 25(2), "
        "383–417.")
    p(doc, "2. Modèles d'évaluation conditionnels et marche aléatoire",
      indent=1)
    ref(doc, "Samuelson, P. A. (1965). Proof that properly anticipated "
        "prices fluctuate randomly. Industrial Management Review, 6(2), "
        "41–49.")
    p(doc, "3. Le modèle d'évaluation des actifs financiers (MEDAF / CAPM)",
      indent=1)
    ref(doc, "Sharpe, W. F. (1964). Capital asset prices: A theory of "
        "market equilibrium under conditions of risk. The Journal of "
        "Finance, 19(3), 425–442.")
    ref(doc, "Lintner, J. (1965). The valuation of risk assets and the "
        "selection of risky investments. Review of Economics and "
        "Statistics, 47(1), 13–37.")

    p_mixed(doc, [("B. Les anomalies empiriques et la remise en cause "
                   "de l'EMH", True, False)])
    p(doc, "1. Anomalies calendaires, momentum, value premium", indent=1)
    ref(doc, "Jegadeesh, N., & Titman, S. (1993). Returns to buying "
        "winners and selling losers. The Journal of Finance, 48(1), "
        "65–91.")
    ref(doc, "Fama, E. F., & French, K. R. (1993). Common risk factors "
        "in the returns on stocks and bonds. Journal of Financial "
        "Economics, 33(1), 3–56.")
    p(doc, "2. Sous-réaction et sur-réaction aux annonces", indent=1)
    ref(doc, "De Bondt, W. F. M., & Thaler, R. (1985). Does the stock "
        "market overreact? The Journal of Finance, 40(3), 793–805.")

    add_heading_styled(doc, "II. La finance comportementale", level=3)

    p_mixed(doc, [("A. Biais cognitifs et heuristiques", True, False)])
    p(doc, "1. Théorie des perspectives", indent=1)
    ref(doc, "Kahneman, D., & Tversky, A. (1979). Prospect theory: An "
        "analysis of decision under risk. Econometrica, 47(2), 263–291.")
    p(doc, "2. Excès de confiance, ancrage, biais de disponibilité",
      indent=1)
    ref(doc, "Barberis, N., & Thaler, R. (2003). A survey of behavioral "
        "finance. Handbook of the Economics of Finance, 1, 1053–1128.")

    p_mixed(doc, [("B. Implications pour la modélisation algorithmique",
                   True, False)])
    p(doc, "1. Sentiment du marché et bruits informationnels", indent=1)
    ref(doc, "Baker, M., & Wurgler, J. (2007). Investor sentiment in the "
        "stock market. Journal of Economic Perspectives, 21(2), 129–151.")
    p(doc, "2. Justification de l'intégration NLP des actualités", indent=1)

    # Section 2
    add_heading_styled(doc, "Section 2 — Méthodes de valorisation et "
                       "d'analyse", level=2)

    add_heading_styled(doc, "I. L'analyse fondamentale", level=3)

    p_mixed(doc, [("A. Approches par actualisation des flux", True, False)])
    p(doc, "1. Modèle des flux de trésorerie actualisés (DCF)", indent=1)
    ref(doc, "Damodaran, A. (2012). Investment Valuation: Tools and "
        "Techniques for Determining the Value of Any Asset (3rd ed.). "
        "Wiley.")
    p(doc, "2. Modèle de Gordon-Shapiro et Discounted Dividend Model (DDM)",
      indent=1)
    ref(doc, "Gordon, M. J., & Shapiro, E. (1956). Capital equipment "
        "analysis: The required rate of profit. Management Science, "
        "3(1), 102–110.")
    p(doc, "3. Coût moyen pondéré du capital (WACC)", indent=1)
    ref(doc, "Modigliani, F., & Miller, M. H. (1958). The cost of capital, "
        "corporation finance and the theory of investment. The American "
        "Economic Review, 48(3), 261–297.")

    p_mixed(doc, [("B. Approches par les fondamentaux comptables",
                   True, False)])
    p(doc, "1. Méthode de Benjamin Graham (valeur intrinsèque)", indent=1)
    ref(doc, "Graham, B., & Dodd, D. (1934 / 2008 ed.). Security Analysis. "
        "McGraw-Hill.")
    p(doc, "2. Multiples de marché : PER, P/B, P/S, EV/EBITDA", indent=1)
    ref(doc, "Liu, J., Nissim, D., & Thomas, J. (2002). Equity valuation "
        "using multiples. Journal of Accounting Research, 40(1), 135–172.")
    p(doc, "3. Spécificités de la valorisation des banques", indent=1)
    ref(doc, "Damodaran, A. (2009). Valuing financial service firms. "
        "Stern School of Business Working Paper.")

    p_mixed(doc, [("C. Approches stochastiques : la simulation Monte Carlo",
                   True, False)])
    p(doc, "1. Mouvement brownien géométrique (GBM)", indent=1)
    ref(doc, "Black, F., & Scholes, M. (1973). The pricing of options and "
        "corporate liabilities. Journal of Political Economy, 81(3), "
        "637–654.")
    p(doc, "2. Application aux trajectoires de prix", indent=1)
    ref(doc, "Glasserman, P. (2003). Monte Carlo Methods in Financial "
        "Engineering. Springer.")

    add_heading_styled(doc, "II. L'analyse technique", level=3)

    p_mixed(doc, [("A. Hypothèses fondatrices et critique académique",
                   True, False)])
    p(doc, "1. Trois principes de Charles Dow", indent=1)
    ref(doc, "Murphy, J. J. (1999). Technical Analysis of the Financial "
        "Markets. New York Institute of Finance.")
    p(doc, "2. Validation empirique partielle", indent=1)
    ref(doc, "Brock, W., Lakonishok, J., & LeBaron, B. (1992). Simple "
        "technical trading rules and the stochastic properties of stock "
        "returns. The Journal of Finance, 47(5), 1731–1764.")
    ref(doc, "Lo, A. W., Mamaysky, H., & Wang, J. (2000). Foundations of "
        "technical analysis. The Journal of Finance, 55(4), 1705–1765.")

    p_mixed(doc, [("B. Indicateurs techniques retenus", True, False)])
    p(doc, "1. RSI (Wilder, 1978)", indent=1)
    ref(doc, "Wilder, J. W. (1978). New Concepts in Technical Trading "
        "Systems. Trend Research.")
    p(doc, "2. MACD (Appel, 1979)", indent=1)
    ref(doc, "Appel, G. (2005). Technical Analysis: Power Tools for "
        "Active Investors. Financial Times Prentice Hall.")
    p(doc, "3. Bandes de Bollinger", indent=1)
    ref(doc, "Bollinger, J. (2001). Bollinger on Bollinger Bands. "
        "McGraw-Hill.")
    p(doc, "4. Average True Range (ATR) et gestion du risque", indent=1)
    ref(doc, "Wilder, J. W. (1978). New Concepts in Technical Trading "
        "Systems. Trend Research.")

    p_mixed(doc, [("C. Microstructure du marché et carnet d'ordres",
                   True, False)])
    p(doc, "1. Order imbalance et pression directionnelle", indent=1)
    ref(doc, "Cont, R., Kukanov, A., & Stoikov, S. (2014). The price "
        "impact of order book events. Journal of Financial Econometrics, "
        "12(1), 47–88.")
    p(doc, "2. Spread bid-ask et liquidité", indent=1)
    ref(doc, "Amihud, Y. (2002). Illiquidity and stock returns. Journal "
        "of Financial Markets, 5(1), 31–56.")

    # Section 3
    add_heading_styled(doc, "Section 3 — Trading algorithmique et "
                       "systèmes multi-signaux", level=2)

    add_heading_styled(doc, "I. Trading algorithmique", level=3)
    p_mixed(doc, [("A. Définitions, typologies, historique", True, False)])
    p(doc, "1. Du trading systématique au trading haute fréquence", indent=1)
    ref(doc, "Aldridge, I. (2013). High-Frequency Trading (2nd ed.). Wiley.")
    p(doc, "2. Cadres réglementaires (MiFID II, AMF, AMMC)", indent=1)

    p_mixed(doc, [("B. Apprentissage statistique et IA en finance",
                   True, False)])
    p(doc, "1. Machine learning supervisé en prédiction de rendements",
      indent=1)
    ref(doc, "Gu, S., Kelly, B., & Xiu, D. (2020). Empirical asset pricing "
        "via machine learning. The Review of Financial Studies, 33(5), "
        "2223–2273.")
    p(doc, "2. Limites et risques", indent=1)
    ref(doc, "López de Prado, M. (2018). Advances in Financial Machine "
        "Learning. Wiley.")

    add_heading_styled(doc, "II. Systèmes à base de règles pondérées et "
                       "agents de décision", level=3)
    p_mixed(doc, [("A. Justification du choix non-ML", True, False)])
    p(doc, "1. Auditabilité et transparence", indent=1)
    ref(doc, "Rudin, C. (2019). Stop explaining black box machine learning "
        "models for high stakes decisions. Nature Machine Intelligence, "
        "1(5), 206–215.")
    p(doc, "2. Surapprentissage en données limitées (cas marché émergent)",
      indent=1)

    p_mixed(doc, [("B. Architectures multi-agents et fusion de signaux",
                   True, False)])
    ref(doc, "Wooldridge, M. (2009). An Introduction to MultiAgent Systems "
        "(2nd ed.). Wiley.")
    ref(doc, "Chan, E. P. (2013). Algorithmic Trading: Winning Strategies "
        "and Their Rationale. Wiley.")

    # Section 4
    add_heading_styled(doc, "Section 4 — Traitement automatique du langage "
                       "naturel appliqué à la finance", level=2)

    add_heading_styled(doc, "I. NLP financier : panorama", level=3)
    p_mixed(doc, [("A. Analyse de sentiment textuel", True, False)])
    ref(doc, "Loughran, T., & McDonald, B. (2011). When is a liability "
        "not a liability? Textual analysis, dictionaries, and 10-Ks. "
        "The Journal of Finance, 66(1), 35–65.")
    ref(doc, "Tetlock, P. C. (2007). Giving content to investor sentiment: "
        "The role of media in the stock market. The Journal of Finance, "
        "62(3), 1139–1168.")

    p_mixed(doc, [("B. Modèles de langage et embeddings financiers",
                   True, False)])
    ref(doc, "Araci, D. (2019). FinBERT: Financial sentiment analysis "
        "with pre-trained language models. arXiv:1908.10063.")
    ref(doc, "Yang, Y., Uy, M. C. S., & Huang, A. (2020). FinBERT: A "
        "pretrained language model for financial communications. "
        "arXiv:2006.08097.")

    add_heading_styled(doc, "II. Mesure de l'impact informationnel sur "
                       "les rendements", level=3)
    ref(doc, "Bollen, J., Mao, H., & Zeng, X. (2011). Twitter mood "
        "predicts the stock market. Journal of Computational Science, "
        "2(1), 1–8.")

    # Section 5
    add_heading_styled(doc, "Section 5 — Gestion du risque et évaluation "
                       "de stratégies", level=2)

    add_heading_styled(doc, "I. Mesures de risque", level=3)
    p(doc, "VaR historique et paramétrique", indent=0.5)
    ref(doc, "Jorion, P. (2007). Value at Risk: The New Benchmark for "
        "Managing Financial Risk (3rd ed.). McGraw-Hill.")
    p(doc, "Drawdown maximal et conditional VaR", indent=0.5)

    add_heading_styled(doc, "II. Mesures de performance ajustée au risque",
                       level=3)
    p(doc, "Ratio de Sharpe", indent=0.5)
    ref(doc, "Sharpe, W. F. (1994). The Sharpe ratio. The Journal of "
        "Portfolio Management, 21(1), 49–58.")
    p(doc, "Ratio de Sortino", indent=0.5)
    ref(doc, "Sortino, F. A., & Price, L. N. (1994). Performance "
        "measurement in a downside risk framework. The Journal of "
        "Investing, 3(3), 59–64.")
    p(doc, "Ratio de Calmar et Information Ratio", indent=0.5)

    add_heading_styled(doc, "III. Backtesting : méthodologie et biais",
                       level=3)
    ref(doc, "Bailey, D. H., Borwein, J., López de Prado, M., & Zhu, Q. J. "
        "(2014). Pseudo-mathematics and financial charlatanism: The "
        "effects of backtest overfitting. Notices of the AMS, 61(5), "
        "458–471.")

    # Section 6
    add_heading_styled(doc, "Section 6 — Spécificités des marchés émergents "
                       "et de la Bourse de Casablanca", level=2)

    add_heading_styled(doc, "I. Marchés émergents : caractéristiques et "
                       "anomalies", level=3)
    ref(doc, "Bekaert, G., & Harvey, C. R. (2002). Research in emerging "
        "markets finance: Looking to the future. Emerging Markets Review, "
        "3(4), 429–448.")

    add_heading_styled(doc, "II. Études empiriques sur le marché marocain",
                       level=3)
    p(doc, "Efficience faible de la BVC", indent=0.5)
    ref(doc, "El Khattab, Y., & Moudine, C. (2014). Test de l'efficience "
        "faible du marché boursier marocain. Revue d'Études en Management "
        "et Finance d'Organisation.")
    p(doc, "Volatilité et liquidité du MASI", indent=0.5)
    ref(doc, "Publications AMMC, Bank Al-Maghrib, Conseil Déontologique "
        "des Valeurs Mobilières.")

    p_mixed(doc, [("Conclusion du Chapitre 1", True, True)])
    p(doc, "Synthèse et positionnement épistémologique du travail.",
      italic=True)

    doc.add_page_break()

    # ============================================================
    # CHAPITRE 2
    # ============================================================
    add_heading_styled(
        doc,
        "CHAPITRE 2 — CONTEXTE D'APPLICATION : LE MARCHÉ MAROCAIN, "
        "ATTIJARIWAFA BANK ET PEAQOCK (pages 46–68)", level=1)

    p_mixed(doc, [("Introduction du Chapitre 2", True, True)])

    add_heading_styled(doc, "Section 1 — Le marché financier marocain et "
                       "la Bourse de Casablanca", level=2)
    add_heading_styled(doc, "I. Histoire et structure institutionnelle",
                       level=3)
    p(doc, "A. De la création (1929) à la modernisation (1993, 2016, 2021)",
      indent=0.5)
    p(doc, "B. Les acteurs : AMMC, BVC, Maroclear, sociétés de bourse",
      indent=0.5)
    p(doc, "C. Compartiments, indices (MASI, MASI ESG, MADEX) et règles "
      "de cotation", indent=0.5)

    add_heading_styled(doc, "II. Capitalisation, liquidité et profondeur",
                       level=3)
    p(doc, "A. Évolution de la capitalisation boursière (2010–2025)",
      indent=0.5)
    p(doc, "B. Volume moyen quotidien, taux de rotation, free float",
      indent=0.5)
    p(doc, "C. Comparaison régionale (CSE, EGX, BRVM)", indent=0.5)

    add_heading_styled(doc, "III. Cadre réglementaire et fiscalité", level=3)
    p(doc, "A. Loi 19-14 sur la BVC, AMMC, lutte anti-blanchiment",
      indent=0.5)
    p(doc, "B. Régime fiscal des plus-values et dividendes", indent=0.5)

    add_heading_styled(doc, "Section 2 — Attijariwafa Bank : profil et "
                       "position stratégique", level=2)
    add_heading_styled(doc, "I. Présentation de la banque", level=3)
    p(doc, "A. Historique (de la BCM à ATW), gouvernance, actionnariat "
      "(SNI / Al Mada)", indent=0.5)
    p(doc, "B. Lignes métiers et présence internationale (Afrique, Europe)",
      indent=0.5)

    add_heading_styled(doc, "II. Profil boursier et fondamentaux", level=3)
    p(doc, "A. Évolution du cours, capitalisation, poids dans le MASI",
      indent=0.5)
    p(doc, "B. Indicateurs financiers : PER, ROE, dividend yield, P/B "
      "(5 ans)", indent=0.5)
    p(doc, "C. Performance comparée au secteur bancaire marocain (BCP, "
      "BMCE, CIH)", indent=0.5)

    add_heading_styled(doc, "III. Pourquoi ATW comme cas d'étude", level=3)
    p(doc, "A. Liquidité supérieure à la moyenne BVC", indent=0.5)
    p(doc, "B. Couverture analyste plus dense", indent=0.5)
    p(doc, "C. Sensibilité aux variables macroéconomiques marocaines",
      indent=0.5)

    add_heading_styled(doc, "Section 3 — L'entreprise Peaqock et le "
                       "périmètre du projet", level=2)
    add_heading_styled(doc, "I. Présentation de Peaqock", level=3)
    p(doc, "A. Mission, expertises (data, IA appliquée à la finance et "
      "l'industrie)", indent=0.5)
    p(doc, "B. Positionnement marché et clients", indent=0.5)

    add_heading_styled(doc, "II. Genèse et périmètre du PFE", level=3)
    p(doc, "A. Cadrage initial du besoin", indent=0.5)
    p(doc, "B. Livrables attendus", indent=0.5)
    p(doc, "C. Calendrier et méthodologie de pilotage", indent=0.5)

    p_mixed(doc, [("Conclusion du Chapitre 2", True, True)])

    doc.add_page_break()

    # ============================================================
    # CHAPITRE 3
    # ============================================================
    add_heading_styled(
        doc,
        "CHAPITRE 3 — CONCEPTION ET IMPLÉMENTATION DE L'AGENT "
        "MULTI-SIGNAUX (pages 69–110)", level=1)

    p_mixed(doc, [("Introduction du Chapitre 3", True, True)])

    add_heading_styled(doc, "Section 1 — Conception de la solution", level=2)
    add_heading_styled(doc, "I. Architecture globale", level=3)
    p(doc, "A. Vision macro : pipeline en 4 couches (collecte → stockage "
      "→ modélisation → synthèse)", indent=0.5)
    p(doc, "B. Diagramme d'architecture général (figure)", indent=0.5)
    p(doc, "C. Choix technologiques justifiés", indent=0.5)
    p(doc, "1. Python pour la flexibilité scientifique", indent=1)
    p(doc, "2. PostgreSQL conteneurisé sous Docker pour l'idempotence "
      "et la portabilité", indent=1)
    p(doc, "3. Architecture modulaire : scrapers/, news_crawler/, models/, "
      "agents/, database/, autorun/", indent=1)

    add_heading_styled(doc, "II. Pipeline de collecte des données", level=3)
    p(doc, "A. Scraper marché temps réel (Medias24 + API BVC) : tables "
      "bourse_daily, bourse_intraday, bourse_orderbook, "
      "technicals_snapshot", indent=0.5)
    p(doc, "B. Scraper macroéconomique (Banque Mondiale, FMI, yfinance, "
      "Investing) : table macro_morocco", indent=0.5)
    p(doc, "C. Scraper fondamentaux (MarketScreener via Selenium) : "
      "tables fondamental_snapshot, fondamental_yearly", indent=0.5)
    p(doc, "D. Crawlers d'actualités : Boursenews, Médias24, L'Économiste, "
      "Aujourd'hui le Maroc, MarketScreener, Google News", indent=0.5)
    p(doc, "E. Logique de canonicalisation, déduplication et signal "
      "scoring (0–100)", indent=0.5)

    add_heading_styled(doc, "III. Architecture de l'agent", level=3)
    p(doc, "A. Modèle conceptuel (chaque signal → score normalisé → "
      "pondération → verdict)", indent=0.5)
    p(doc, "B. Schéma logique de la fusion des signaux", indent=0.5)
    p(doc, "C. Choix d'un système à règles pondérées plutôt qu'un modèle "
      "ML (transparence, données limitées, robustesse)", indent=0.5)

    add_heading_styled(doc, "Section 2 — Modélisation des signaux", level=2)

    add_heading_styled(doc, "I. Module fondamental", level=3)
    p(doc, "A. Calcul des ratios : PER, P/B, ROE, dividend yield, "
      "croissance des bénéfices", indent=0.5)
    p(doc, "B. Modèle DCF : projection FCF, terminal value, WACC bancaire "
      "ajusté", indent=0.5)
    p(doc, "C. Modèle DDM (Gordon-Shapiro) adapté au profil de "
      "distribution d'ATW", indent=0.5)
    p(doc, "D. Modèle de Graham : valeur intrinsèque conservatrice",
      indent=0.5)
    p(doc, "E. Valorisation relative par multiples sectoriels", indent=0.5)
    p(doc, "F. Agrégation : moyenne (et discussion d'une pondération par "
      "précision historique)", indent=0.5)

    add_heading_styled(doc, "II. Module technique", level=3)
    p(doc, "A. ATR(14) selon la formule de Wilder : volatilité et "
      "calibrage des cibles de trade", indent=0.5)
    p(doc, "B. Volatilité 20 jours et conversion correcte (lien avec le "
      "bug identifié dans ag.py)", indent=0.5)
    p(doc, "C. RSI, MACD, Bollinger Bands (lus depuis "
      "technicals_snapshot)", indent=0.5)
    p(doc, "D. Features carnet d'ordres : Order Imbalance multi-niveaux, "
      "spread bid-ask en bps, Z-score intra-séance, profondeur, VWMP",
      indent=0.5)

    add_heading_styled(doc, "III. Module macroéconomique", level=3)
    p(doc, "A. Variables retenues : croissance PIB, inflation CPI, dette "
      "publique/PIB, USD/MAD, M2", indent=0.5)
    p(doc, "B. Garde-fous et bug IMF DataMapper (sanity bands, fallback "
      "Banque Mondiale)", indent=0.5)
    p(doc, "C. Score macro composite", indent=0.5)

    add_heading_styled(doc, "IV. Module NLP / sentiment", level=3)
    p(doc, "A. Périmètre : 6 sources marocaines, fenêtre glissante "
      "14 jours", indent=0.5)
    p(doc, "B. Pipeline texte : nettoyage, tokenisation, normalisation",
      indent=0.5)
    p(doc, "C. Calcul du signal_score (0–100) et flag is_atw_core",
      indent=0.5)
    p(doc, "D. Discussion : sentiment dérivé du score d'article ; piste "
      "d'amélioration via FinBERT/lexique Loughran-McDonald traduit",
      indent=0.5)

    add_heading_styled(doc, "V. Détection du régime de marché", level=3)
    p(doc, "A. Logique : moyenne pondérée des rendements 1S/1M/3M, "
      "position vs MA20, volatilité", indent=0.5)
    p(doc, "B. Trois régimes : BULL / BEAR / SIDEWAYS", indent=0.5)
    p(doc, "C. Confiance associée", indent=0.5)

    add_heading_styled(doc, "Section 3 — Mécanisme de pondération et "
                       "logique de décision", level=2)

    add_heading_styled(doc, "I. Prédiction trading court terme (4 semaines)",
                       level=3)
    p(doc, "A. Cibles dérivées de l'ATR : zone d'entrée, objectif "
      "(1.5×ATR), stop (2.0×ATR)", indent=0.5)
    p(doc, "B. Simulation Monte Carlo GBM (1000 trajectoires sur 20 jours)",
      indent=0.5)
    p(doc, "C. Probabilité de gain, VaR 95%, IC 90%", indent=0.5)
    p(doc, "D. Confiance HIGH/MEDIUM/LOW", indent=0.5)

    add_heading_styled(doc, "II. Prédiction investissement long terme "
                       "(12 mois)", level=3)
    p(doc, "A. Juste valeur agrégée et upside", indent=0.5)
    p(doc, "B. Seuils dynamiques adaptatifs (volatilité × multiplicateur "
      "de régime × facteur d'incertitude)", indent=0.5)
    p(doc, "C. Verdict ACHAT / CONSERVER / VENDRE", indent=0.5)

    add_heading_styled(doc, "III. Synthèse LLM avec validation de citations",
                       level=3)
    p(doc, "A. Bloc d'evidence horodaté (identifiants [MKT-*], [VAL-*], "
      "[NEWS-*], etc.)", indent=0.5)
    p(doc, "B. Appel LLM Groq avec contrat JSON strict", indent=0.5)
    p(doc, "C. Validation regex des citations + fallback déterministe si "
      "échec", indent=0.5)

    add_heading_styled(doc, "Section 4 — Backtesting et évaluation", level=2)

    add_heading_styled(doc, "I. Méthodologie", level=3)
    p(doc, "A. Période d'évaluation, walk-forward analysis", indent=0.5)
    p(doc, "B. Benchmarks : buy-and-hold ATW, indice MASI", indent=0.5)
    p(doc, "C. Frais de transaction et slippage hypothétiques", indent=0.5)

    add_heading_styled(doc, "II. Métriques de performance", level=3)
    p(doc, "A. Rendement cumulé, rendement annualisé", indent=0.5)
    p(doc, "B. Ratios de Sharpe, Sortino, Calmar", indent=0.5)
    p(doc, "C. Drawdown maximal et durée de drawdown", indent=0.5)
    p(doc, "D. Hit ratio et profit factor", indent=0.5)

    add_heading_styled(doc, "III. Résultats", level=3)
    p(doc, "A. Tableaux comparatifs et courbes d'équité", indent=0.5)
    p(doc, "B. Décomposition de la contribution par signal", indent=0.5)
    p(doc, "C. Robustesse : analyse de sensibilité aux paramètres clés "
      "(seuils, pondérations)", indent=0.5)

    add_heading_styled(doc, "Section 5 — Limites identifiées et "
                       "recommandations", level=2)

    add_heading_styled(doc, "I. Limites techniques constatées dans l'agent",
                       level=3)
    limites_tech = [
        "Bug de double conversion de volatilité dans la simulation "
        "Monte Carlo",
        "Drift Monte Carlo biaisé par le mélange ret_1w / ret_1m",
        "np.random.seed(42) figeant la stochasticité",
        "Variables macro câblées à None (chômage, taux directeur)",
        "Sentiment dérivé du signal_score plutôt que d'une vraie analyse "
        "linguistique",
        "Baseline de volatilité 2.5 % codée en dur, non calibrée pour "
        "Casa-Bourse",
        "Instabilité du LLM (dépendance externe Groq)",
    ]
    for it in limites_tech:
        doc.add_paragraph(it, style="List Bullet")

    add_heading_styled(doc, "II. Limites méthodologiques", level=3)
    limites_met = [
        "Mono-titre (ATW seul, pas de comparaison sectorielle)",
        "Hypothèses GBM (rendements log-normaux i.i.d.)",
        "Trois régimes seulement (manque crash, recovery)",
        "Risque de surapprentissage des seuils sur l'historique de backtest",
    ]
    for it in limites_met:
        doc.add_paragraph(it, style="List Bullet")

    add_heading_styled(doc, "III. Recommandations d'extension", level=3)
    recos = [
        ("A. ", "Calibrage des seuils par optimisation bayésienne sur "
         "historique élargi"),
        ("B. ", "Intégration d'un modèle FinBERT bilingue (français/arabe) "
         "pour le sentiment"),
        ("C. ", "Élargissement à un panier de titres BVC (BCP, BMCE, IAM) "
         "avec corrélations sectorielles"),
        ("D. ", "Régimes augmentés (5 états avec HMM) et volatilité "
         "implicite si options disponibles"),
        ("E. ", "Pondération adaptative des modèles de valorisation par "
         "leur précision out-of-sample"),
    ]
    for label, txt in recos:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5)
        r1 = para.add_run(label)
        r1.bold = True
        r1.font.name = "Calibri"
        r2 = para.add_run(txt)
        r2.font.name = "Calibri"

    p_mixed(doc, [("Conclusion du Chapitre 3", True, True)])

    doc.add_page_break()

    # ============================================================
    # CONCLUSION GÉNÉRALE
    # ============================================================
    add_heading_styled(doc,
                       "CONCLUSION GÉNÉRALE (pages 111–115)", level=1)

    add_heading_styled(doc, "1. Synthèse des résultats par chapitre",
                       level=2)
    syntheses = [
        "Chapitre 1 : ancrage théorique multi-disciplinaire",
        "Chapitre 2 : caractérisation du marché marocain et du cas ATW",
        "Chapitre 3 : conception et évaluation de l'agent",
    ]
    for s in syntheses:
        doc.add_paragraph(s, style="List Bullet")

    add_heading_styled(doc, "2. Réponse à la problématique", level=2)
    p(doc, "Un agent multi-signaux à règles pondérées peut effectivement "
      "améliorer la qualité de la décision sur la BVC, sous réserve d'une "
      "discipline méthodologique stricte (validation, backtesting, "
      "calibrage).")

    add_heading_styled(doc, "3. Contributions", level=2)
    p_mixed(doc, [
        ("Académiques : ", True, True),
        ("enrichissement de la littérature sur les marchés émergents "
         "francophones et sur les architectures hybrides quantitatif/NLP.",
         False, False)
    ])
    p_mixed(doc, [
        ("Professionnelles : ", True, True),
        ("livrable opérationnel pour Peaqock, prototype reproductible et "
         "documenté.", False, False)
    ])

    add_heading_styled(doc, "4. Limites du travail", level=2)
    p(doc, "Rappel synthétique des limites techniques et méthodologiques.",
      italic=True)

    add_heading_styled(doc, "5. Perspectives", level=2)
    p(doc, "Extension multi-titres, temps réel, dashboard mobile, "
      "intégration des critères ESG.")

    doc.add_page_break()

    # ============================================================
    # BIBLIO + ANNEXES
    # ============================================================
    add_heading_styled(doc, "BIBLIOGRAPHIE (format APA 7e édition)", level=1)
    p(doc, "Liste finale agrégée des références citées dans le mémoire, "
      "organisée alphabétiquement.", italic=True)

    add_heading_styled(doc, "TABLE DES MATIÈRES", level=1)
    p(doc, "(Générée automatiquement en fin de rédaction.)", italic=True)

    add_heading_styled(doc, "ANNEXES", level=1)
    annexes = [
        "Annexe A — Schéma de la base de données PostgreSQL "
        "(DDL 01_schema.sql)",
        "Annexe B — Extraits de code commentés (modules clés de ag.py, "
        "fundamental_models.py)",
        "Annexe C — Captures d'écran de l'agent en exécution et logs",
        "Annexe D — Tableaux détaillés du backtesting",
        "Annexe E — Liste exhaustive des 6 sources d'actualités et "
        "règles de scoring",
        "Annexe F — Glossaire technique (200+ termes)",
    ]
    for a in annexes:
        doc.add_paragraph(a, style="List Bullet")

    # ============================================================
    # SAVE
    # ============================================================
    out = r"C:\Users\arhou\PFE.01\Plan_Memoire_PFE_Arhoune.docx"
    doc.save(out)
    print(f"Document saved: {out}")


if __name__ == "__main__":
    main()
